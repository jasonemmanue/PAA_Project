"""Routeur `/rapport` — endpoints alignés sur la méthodologie DEESP/DEEF.

Reproduit les 17 tableaux + 12 graphiques du rapport officiel
*« Evaluation du temps de traversée — octobre 2025 »* du PAA.

Endpoints :
  - GET /rapport/temps-theoriques       → Tableau 1
  - GET /rapport/temps-traversee?campagne=AAAA-MM → Tableaux 3-17
  - GET /rapport/zones-congestionnees?campagne=AAAA-MM → Tableau 16
  - GET /rapport/graphique/{troncon_id}?agregat=min|max&campagne=AAAA-MM
                                        → données pour BarChart 1-12
  - GET /rapport/comparaison?campagne_a=AAAA-MM&campagne_b=AAAA-MM → Tableau 19

Le paramètre `campagne` accepte un format `AAAA-MM` (ex. `2026-02`) et est
résolu en (1er du mois, dernier jour du mois) côté serveur.
"""

from __future__ import annotations

import logging
import re
from datetime import date as DateType, datetime, time, timezone
from typing import Any
from zoneinfo import ZoneInfo

from fastapi import APIRouter, Depends, File, HTTPException, Query, Response, UploadFile, status
from sqlalchemy.orm import Session

from app.analyse import rapport_paa
from app.core.config import get_settings
from app.db.session import get_db


# Logger dédié — visible dans Railway sous le tag "paa.rapport"
logger = logging.getLogger("paa.rapport")

router = APIRouter(prefix="/rapport", tags=["rapport DEESP"])


# ---------------------------------------------------------------------------
# Helper — sanitisation pour fpdf2 (Helvetica ne supporte que Latin-1)
# ---------------------------------------------------------------------------


def _sanitize_pdf(texte: str) -> str:
    """Remplace les caractères hors Latin-1 avant écriture dans fpdf2.

    fpdf2 avec la police Helvetica n'accepte que le jeu de caractères
    Latin-1 (ISO 8859-1). Cette fonction convertit les caractères Unicode
    courants en équivalents ASCII, puis encode/décode en Latin-1 (errors=replace)
    pour neutraliser tout résidu.
    """
    remplacements = {
        "→": "->",   "←": "<-",   "↔": "<->",
        "≥": ">=",   "≤": "<=",   "×": "x",    "÷": "/",
        "–": "-",    "—": "-",    "…": "...",
        "’": "'",  "‘": "'",
        "“": '"',  "”": '"',
        "°": " deg",
    }
    for char, repl in remplacements.items():
        texte = texte.replace(char, repl)
    return texte.encode("latin-1", errors="replace").decode("latin-1")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


_REGEX_CAMPAGNE = re.compile(r"^(\d{4})-(\d{2})$")


def _parser_campagne(libelle: str) -> tuple[datetime, datetime]:
    """Convertit 'AAAA-MM' en (debut_utc, fin_utc) couvrant le mois entier."""
    m = _REGEX_CAMPAGNE.match(libelle)
    if not m:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Format attendu : 'AAAA-MM' (ex. '2026-02').",
        )
    annee, mois = int(m.group(1)), int(m.group(2))
    if not (1 <= mois <= 12):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Le mois doit être compris entre 1 et 12.",
        )
    debut, fin = rapport_paa.fenetre_campagne(annee, mois)
    fuseau_local = ZoneInfo(get_settings().tz)
    debut_utc = datetime.combine(debut, time.min, tzinfo=fuseau_local).astimezone(timezone.utc)
    fin_utc = datetime.combine(fin, time.max, tzinfo=fuseau_local).astimezone(timezone.utc)
    return debut_utc, fin_utc


def _bornes_utc(
    campagne: str,
    debut_override: DateType | None,
    fin_override: DateType | None,
) -> tuple[datetime, datetime]:
    """Résout les bornes UTC en tenant compte des overrides optionnels debut/fin."""
    debut_utc, fin_utc = _parser_campagne(campagne)
    fuseau_local = ZoneInfo(get_settings().tz)
    if debut_override is not None:
        debut_utc = datetime.combine(debut_override, time.min, tzinfo=fuseau_local).astimezone(timezone.utc)
    if fin_override is not None:
        fin_utc = datetime.combine(fin_override, time.max, tzinfo=fuseau_local).astimezone(timezone.utc)
    return debut_utc, fin_utc


# ---------------------------------------------------------------------------
# GET /rapport/temps-theoriques
# ---------------------------------------------------------------------------


@router.get(
    "/temps-theoriques",
    summary="Tableau 1 — Temps théoriques 50 km/h par axe",
    description=(
        "Renvoie le Tableau 1 du rapport : distance et temps théorique à "
        "50 km/h pour chacun des 3 axes officiels. Données statiques "
        "dérivées du seed des tronçons."
    ),
)
async def get_temps_theoriques(
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    lignes = rapport_paa.temps_theoriques(db)
    return {
        "tableau": "Tableau 1 — Temps de traversée normal pour 50 km/h",
        "lignes": [
            {
                "axe": tt.axe,
                "distance_km": tt.distance_km,
                "temps_50kmh_s": tt.temps_50kmh_s,
                "temps_50kmh": tt.temps_50kmh_str,
            }
            for tt in lignes
        ],
    }


# ---------------------------------------------------------------------------
# GET /rapport/temps-traversee
# ---------------------------------------------------------------------------


@router.get(
    "/temps-traversee",
    summary="Tableaux 3-17 — Min/Moyen/Max par tronçon × type-jour",
    description=(
        "Renvoie pour chaque tronçon et chaque type de jour "
        "(`jour_ouvrable` / `week_end`) les temps minimal, moyen et maximal "
        "en minutes sur la campagne demandée. C'est la base de tous les "
        "tableaux 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15 et 17."
    ),
)
async def get_temps_traversee(
    campagne: str = Query(
        ..., description="Format 'AAAA-MM', ex. '2026-02' pour février 2026",
    ),
    debut: DateType | None = Query(None, description="Date de début (YYYY-MM-DD) — affine la fenêtre à l'intérieur du mois."),
    fin: DateType | None = Query(None, description="Date de fin (YYYY-MM-DD) — affine la fenêtre à l'intérieur du mois."),
    heure_debut: int = Query(0, ge=0, le=23, description="Heure de début de la plage (incluse, 0-23). Défaut 0 = 24h/24."),
    heure_fin: int = Query(24, ge=1, le=24, description="Heure de fin de la plage (exclue, 1-24). Défaut 24 = 24h/24."),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    debut_utc, fin_utc = _bornes_utc(campagne, debut, fin)
    stats = rapport_paa.temps_traversee_par_troncon(db, debut_utc, fin_utc, heure_debut=heure_debut, heure_fin=heure_fin)
    return {
        "campagne": campagne,
        "debut_utc": debut_utc.isoformat(),
        "fin_utc": fin_utc.isoformat(),
        "nb_lignes": len(stats),
        "lignes": [
            {
                "troncon_id": s.troncon_id,
                "troncon_nom": s.troncon_nom,
                "type_jour": s.type_jour,
                "nb_mesures": s.nb_mesures,
                "temps_min_mn": s.temps_min_mn,
                "temps_moyen_mn": s.temps_moyen_mn,
                "temps_max_mn": s.temps_max_mn,
            }
            for s in stats
        ],
    }


# ---------------------------------------------------------------------------
# GET /rapport/zones-congestionnees
# ---------------------------------------------------------------------------


@router.get(
    "/zones-congestionnees",
    summary="Tableau 16 — Tronçons congestionnés selon règles DEESP",
    description=(
        "Applique les règles du § 4.5.3 du CLAUDE.md (extraites du rapport) :\n\n"
        "- **Règle JOUR** : tronçon congestionné si ≥ 3 fois sur un même "
        "  jour-indicatif (ex. 3 lundis sur 4) à la même heure.\n"
        "- **Règle SEMAINE** : tronçon congestionné si ≥ 4 fois à la même "
        "  heure dans la semaine, tous jours confondus.\n\n"
        "Le critère de congestion d'une mesure individuelle est "
        "Couleur Google Maps lue par tronçon : ROUGE OU ORANGE ≥ 50 % → "
        "congestionné (cf. rapport DEESP/DEEF oct. 2025, section "
        "METHODOLOGIE)."
    ),
)
async def get_zones_congestionnees(
    campagne: str = Query(..., description="Format 'AAAA-MM'."),
    debut: DateType | None = Query(None, description="Date de début (YYYY-MM-DD) — affine la fenêtre à l'intérieur du mois."),
    fin: DateType | None = Query(None, description="Date de fin (YYYY-MM-DD) — affine la fenêtre à l'intérieur du mois."),
    heure_debut: int = Query(0, ge=0, le=23, description="Heure de début de la plage (incluse, 0-23). Défaut 0 = 24h/24."),
    heure_fin: int = Query(24, ge=1, le=24, description="Heure de fin de la plage (exclue, 1-24). Défaut 24 = 24h/24."),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    debut_utc, fin_utc = _bornes_utc(campagne, debut, fin)
    cong = rapport_paa.troncons_congestionnes(db, debut_utc, fin_utc, heure_debut=heure_debut, heure_fin=heure_fin)
    nb_jours = max(1, (fin_utc - debut_utc).days + 1)
    seuil_j, seuil_s = rapport_paa.seuils_congestion(debut_utc, fin_utc)
    return {
        "campagne": campagne,
        "nb_jours_plage": nb_jours,
        "nb_entrees": len(cong),
        "regles": {
            "critere_mesure": (
                "Couleur Google Maps : ROUGE OU ORANGE sur ≥ 50 % du tronçon"
            ),
            "seuil_orange_long_pct": 50.0,
            "seuil_jour_effectif": seuil_j,
            "seuil_semaine_effectif": seuil_s,
            "regle_jour_indicatif": f"≥ {seuil_j} occurrence(s) sur le même jour de la semaine",
            "regle_semaine": f"≥ {seuil_s} occurrence(s) à la même heure dans la semaine",
            "adaptatif": nb_jours < 28,
        },
        "entrees": [
            {
                "troncon_id": c.troncon_id,
                "troncon_nom": c.troncon_nom,
                "sous_troncon_id": c.sous_troncon_id,
                "sous_troncon_code": c.sous_troncon_code,
                "sous_troncon_nom": c.sous_troncon_nom,
                "heure": c.heure,
                "tranche": f"{c.heure:02d}h-{c.heure + 1:02d}h",
                "nb_par_jour_semaine": c.nb_jours_congestionnes_par_type,
                "nb_total_semaine": c.nb_total_semaine,
                "regle_jour_indicatif": c.regle_jour_indicatif,
                "regle_semaine": c.regle_semaine,
            }
            for c in cong
        ],
    }


# ---------------------------------------------------------------------------
# GET /rapport/zones-congestionnees/pdf — téléchargement direct
# ---------------------------------------------------------------------------


@router.get(
    "/zones-congestionnees/pdf",
    summary="Tableau 16 au format PDF (téléchargement direct)",
    description=(
        "Génère un PDF A4 paysage contenant le Tableau 16. Le navigateur "
        "déclenchera un téléchargement immédiat grâce à l'en-tête "
        "Content-Disposition: attachment."
    ),
    response_class=Response,
)
async def get_zones_congestionnees_pdf(
    campagne: str = Query(..., description="Format 'AAAA-MM'."),
    debut: DateType | None = Query(None),
    fin: DateType | None = Query(None),
    heure_debut: int = Query(0, ge=0, le=23),
    heure_fin: int = Query(24, ge=1, le=24),
    db: Session = Depends(get_db),
) -> Response:
    logger.info(
        "GET /rapport/zones-congestionnees/pdf — campagne=%r debut=%s fin=%s heure=%d-%d",
        campagne, debut, fin, heure_debut, heure_fin,
    )
    try:
        from fpdf import FPDF  # import local : évite de charger fpdf au démarrage
        logger.info("fpdf2 importé avec succès")
    except ImportError as exc:
        logger.error("fpdf2 manquant : %s", exc)
        raise HTTPException(status_code=500, detail=f"fpdf2 non disponible : {exc}")

    debut_utc, fin_utc = _bornes_utc(campagne, debut, fin)
    logger.info("Bornes calculées : debut_utc=%s fin_utc=%s", debut_utc, fin_utc)

    try:
        cong = rapport_paa.troncons_congestionnes(db, debut_utc, fin_utc, heure_debut=heure_debut, heure_fin=heure_fin)
        logger.info("Données récupérées : %d zone(s) congestionnée(s)", len(cong))
    except Exception as exc:
        logger.error("Erreur récupération données congestion : %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Erreur données : {exc}")

    seuil_j, seuil_s = rapport_paa.seuils_congestion(debut_utc, fin_utc)
    nb_jours = max(1, (fin_utc - debut_utc).days + 1)

    pdf = FPDF(orientation="L", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=10)
    pdf.add_page()

    # En-tête
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 7, "Tableau 16 - Axes congestionnes (regles DEESP)", ln=True)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(26, 54, 93)
    pdf.cell(0, 5, f"Campagne : {campagne}    -    {nb_jours} jour(s) analyse(s)", ln=True)
    pdf.set_text_color(85, 85, 85)
    pdf.set_font("Helvetica", "", 8)
    desc = (
        f"Critere par mesure : couleur Google Maps - ROUGE OU ORANGE >= 50%. "
        f"Seuils appliques : >= {seuil_j} occurrence(s) / jour-indicatif OU "
        f">= {seuil_s} occurrence(s) / semaine."
    )
    pdf.multi_cell(0, 4, desc)
    pdf.ln(2)

    # En-tête du tableau
    pdf.set_fill_color(26, 54, 93)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 8)
    largeurs = [60, 35, 25, 18, 30, 110]  # mm — total = 278
    entetes = ["AXE", "TRONCON CODIFIE", "TRANCHE", "NB/SEM.", "REGLE", "REPARTITION PAR JOUR"]
    for w, h in zip(largeurs, entetes):
        pdf.cell(w, 7, h, border=1, fill=True, align="L")
    pdf.ln()

    # Lignes
    pdf.set_text_color(17, 17, 17)
    pdf.set_font("Helvetica", "", 8)
    if not cong:
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(107, 114, 128)
        pdf.cell(sum(largeurs), 10, "Aucun troncon congestionne sur cette campagne.",
                 border=1, align="C")
    else:
        for i, c in enumerate(cong):
            if i % 2 == 1:
                pdf.set_fill_color(249, 250, 251)
                fill = True
            else:
                pdf.set_fill_color(255, 255, 255)
                fill = True
            sous = (
                f"{c.sous_troncon_code} - {c.sous_troncon_nom or ''}"
                if c.sous_troncon_code
                else "axe entier"
            )
            tranche = f"{c.heure:02d}h-{c.heure + 1:02d}h"
            regles = []
            if c.regle_jour_indicatif:
                regles.append(f">={seuil_j}/jour")
            if c.regle_semaine:
                regles.append(f">={seuil_s}/sem")
            regle_txt = " | ".join(regles) or "-"
            repartition = " ".join(
                f"{j[:3]}:{n}" for j, n in c.nb_jours_congestionnes_par_type.items()
            )

            ligne = [
                (largeurs[0], _sanitize_pdf((c.troncon_nom or "")[:45])),
                (largeurs[1], _sanitize_pdf(sous[:25])),
                (largeurs[2], tranche),
                (largeurs[3], str(c.nb_total_semaine)),
                (largeurs[4], regle_txt),
                (largeurs[5], _sanitize_pdf(repartition[:80])),
            ]
            for w, txt in ligne:
                pdf.cell(w, 6, txt, border=1, fill=fill, align="L")
            pdf.ln()

    # Pied de page
    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(107, 114, 128)
    pdf.cell(0, 4,
             "Genere par FLUIDIS - Port Autonome d'Abidjan - "
             f"Methodologie DEESP rapport octobre 2025", ln=True)

    # Réponse — bytearray → bytes pour httpx/starlette
    try:
        contenu = bytes(pdf.output())
        logger.info("PDF généré avec succès : %d octets", len(contenu))
    except Exception as exc:
        logger.error("Erreur génération PDF fpdf2 : %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Erreur génération PDF : {exc}")

    nom = f"tableau16_{campagne}.pdf"
    return Response(
        content=contenu,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{nom}"'},
    )


# ---------------------------------------------------------------------------
# GET /rapport/matrice-congestion — matrice brute heure × date par tronçon
# ---------------------------------------------------------------------------


@router.get(
    "/matrice-congestion",
    summary="Matrice de congestion par créneau horaire × date",
    description=(
        "Pour le tronçon sélectionné et la plage de dates, renvoie pour chaque "
        "créneau horaire DEESP (07h-19h) et chaque date la mesure Google la plus "
        "récente : congestionné / fluide / indéterminé / pas de mesure. "
        "Permet de visualiser les patterns de congestion sans agrégation."
    ),
)
async def get_matrice_congestion(
    campagne: str = Query(..., description="Format 'AAAA-MM'."),
    troncon_id: int = Query(..., description="ID du tronçon à analyser."),
    debut: DateType | None = Query(None),
    fin: DateType | None = Query(None),
    heure_debut: int = Query(0, ge=0, le=23),
    heure_fin: int = Query(24, ge=1, le=24),
    sous_troncon_id: int | None = Query(None),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    from app.models.models import SousTroncon, Troncon, axe_sous_troncons as m2m
    logger.info(
        "GET /rapport/matrice-congestion — troncon_id=%d sous=%s campagne=%r",
        troncon_id, sous_troncon_id, campagne,
    )
    troncon = db.get(Troncon, troncon_id)
    if troncon is None:
        raise HTTPException(status_code=404, detail=f"Tronçon {troncon_id} introuvable.")
    sous = None
    if sous_troncon_id is not None:
        sous = db.get(SousTroncon, sous_troncon_id)
        if sous is None:
            raise HTTPException(status_code=404, detail="Sous-tronçon introuvable.")
        # Vérifier parent principal OU rattachement M2M (multi-parent)
        if sous.troncon_id != troncon_id:
            from sqlalchemy import select as sa_select
            lien = db.execute(
                sa_select(m2m.c.axe_id).where(
                    m2m.c.axe_id == troncon_id,
                    m2m.c.sous_troncon_id == sous_troncon_id,
                )
            ).first()
            if lien is None:
                raise HTTPException(status_code=404, detail="Sous-tronçon introuvable.")
    debut_utc, fin_utc = _bornes_utc(campagne, debut, fin)
    result = rapport_paa.matrice_congestion(
        db, troncon_id, debut_utc, fin_utc,
        heure_debut=heure_debut, heure_fin=heure_fin,
        sous_troncon_id=sous_troncon_id,
    )
    nom_affichage = (
        f"{troncon.nom} : {sous.nom_court} ({sous.code})"
        if sous is not None else troncon.nom
    )
    return {"troncon_nom": nom_affichage, **result}


# ---------------------------------------------------------------------------
# GET /rapport/export/word — export complet en .docx (tableaux + graphiques)
# ---------------------------------------------------------------------------


@router.get(
    "/export/word",
    summary="Export Word (.docx) complet du Rapport DEESP",
    description=(
        "Génère un document Word A4 paysage contenant en temps réel : "
        "Tableau 1 (théoriques), Tableaux 3-15 (min/moyen/max), "
        "Graphiques 1-12 (BarChart embarqués en PNG), Tableau 16 "
        "(zones congestionnées). Téléchargement direct côté navigateur."
    ),
    response_class=Response,
)
async def export_rapport_word(
    campagne: str = Query(..., description="Format 'AAAA-MM'."),
    debut: DateType | None = Query(None),
    fin: DateType | None = Query(None),
    heure_debut: int = Query(0, ge=0, le=23),
    heure_fin: int = Query(24, ge=1, le=24),
    db: Session = Depends(get_db),
) -> Response:
    logger.info(
        "GET /rapport/export/word — campagne=%r debut=%s fin=%s",
        campagne, debut, fin,
    )

    # Phase 1 : imports lourds (matplotlib + python-docx). Tracé explicitement
    # pour distinguer un échec d'installation côté Railway d'un échec métier.
    try:
        import io
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from docx import Document
        from docx.enum.section import WD_ORIENTATION
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        logger.exception("Echec import matplotlib/python-docx — paquet manquant ?")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Dependance manquante cote serveur : {exc}. "
                   "Verifier que matplotlib et python-docx sont dans requirements.txt.",
        )

    # Phase 2 : résolution des bornes (peut lever 400 si campagne malformée)
    try:
        debut_utc, fin_utc = _bornes_utc(campagne, debut, fin)
        nb_jours = max(1, (fin_utc - debut_utc).days + 1)
        logger.info(
            "Bornes : debut_utc=%s fin_utc=%s nb_jours=%d",
            debut_utc, fin_utc, nb_jours,
        )
    except HTTPException:
        raise
    except Exception:
        logger.exception("Echec resolution bornes campagne=%r debut=%s fin=%s.",
                         campagne, debut, fin)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Bornes de campagne invalides : campagne={campagne!r}",
        )

    # Phase 3 : extraction des données métier
    try:
        theoriques = rapport_paa.temps_theoriques(db)
        logger.info("temps_theoriques OK — %d ligne(s)", len(theoriques))
    except Exception:
        logger.exception("Echec rapport_paa.temps_theoriques.")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Echec calcul Tableau 1 (temps théoriques) — voir logs serveur.",
        )

    # Variable conservée pour compatibilité avec le code en aval
    _theoriques_marker = True  # noqa: F841
    try:
        stats = rapport_paa.temps_traversee_par_troncon(db, debut_utc, fin_utc, heure_debut=heure_debut, heure_fin=heure_fin)
        logger.info("temps_traversee_par_troncon OK — %d troncons", len(stats) if stats else 0)
    except Exception:
        logger.exception("Echec rapport_paa.temps_traversee_par_troncon.")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Echec calcul Tableaux 3-15 (temps de traversee) — voir logs serveur.",
        )

    try:
        cong = rapport_paa.troncons_congestionnes(db, debut_utc, fin_utc, heure_debut=heure_debut, heure_fin=heure_fin)
        seuil_j, seuil_s = rapport_paa.seuils_congestion(debut_utc, fin_utc)
        logger.info("troncons_congestionnes OK — %d entree(s), seuils jour=%s semaine=%s",
                    len(cong) if cong else 0, seuil_j, seuil_s)
    except Exception:
        logger.exception("Echec rapport_paa.troncons_congestionnes/seuils_congestion.")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Echec calcul Tableau 16 (zones congestionnees) — voir logs serveur.",
        )

    # Tronçons actifs (pour graphiques)
    from sqlalchemy import select
    from app.models.models import Troncon
    try:
        troncons = list(
            db.execute(
                select(Troncon).where(Troncon.actif.is_(True)).order_by(Troncon.id)
            ).scalars()
        )
        logger.info("Troncons actifs : %d", len(troncons))
    except Exception:
        logger.exception("Echec lecture des troncons actifs.")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Echec lecture des troncons actifs — voir logs serveur.",
        )

    # ---- Construction du document ----
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    NAVY = RGBColor(0x1A, 0x36, 0x5D)
    GRIS = RGBColor(0x6B, 0x72, 0x80)
    BLANC = RGBColor(0xFF, 0xFF, 0xFF)

    def shade(cell, hex_color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def ajouter_titre(texte: str, niveau: int = 1) -> None:
        p = doc.add_paragraph()
        run = p.add_run(texte)
        run.bold = True
        run.font.color.rgb = NAVY
        run.font.size = Pt(14 if niveau == 1 else 12)

    def ajouter_paragraphe(texte: str, italique: bool = False, taille: int = 9) -> None:
        p = doc.add_paragraph()
        run = p.add_run(texte)
        run.italic = italique
        run.font.size = Pt(taille)
        run.font.color.rgb = GRIS

    # En-tête
    titre = doc.add_paragraph()
    r = titre.add_run("RAPPORT DEESP — ÉVALUATION DU TEMPS DE TRAVERSÉE")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sous = doc.add_paragraph()
    rs = sous.add_run(f"Campagne : {campagne}    •    {nb_jours} jour(s) analysé(s)")
    rs.font.size = Pt(10)
    rs.font.color.rgb = GRIS
    sous.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ajouter_paragraphe(
        "Port Autonome d'Abidjan — Méthodologie DEESP/DEEF "
        "(rapport octobre 2025). Reproduction fidèle des 17 tableaux et "
        "12 graphiques officiels. Données collectées en temps réel via Google Routes.",
        italique=True,
    )
    doc.add_paragraph()

    # ---- Tableau 1 — Temps théoriques 50 km/h ----
    ajouter_titre("Tableau 1 — Temps théoriques à 50 km/h", niveau=1)
    tab1 = doc.add_table(rows=1, cols=3)
    tab1.style = "Light Grid Accent 1"
    hdr = tab1.rows[0].cells
    for i, txt in enumerate(("AXE", "DISTANCE (km)", "TEMPS À 50 km/h")):
        hdr[i].text = txt
        shade(hdr[i], "1A365D")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = BLANC
                run.font.size = Pt(9)
    for tt in theoriques:
        row = tab1.add_row().cells
        row[0].text = tt.axe
        row[1].text = f"{tt.distance_km:.2f}"
        row[2].text = tt.temps_50kmh_str
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    # ---- Tableaux Min / Moyen / Max ----
    libelles = {
        "min": ("Tableaux 3-7 — Temps MINIMAL observé", "temps_min_mn"),
        "moyen": ("Tableaux 8-11 — Temps MOYEN observé", "temps_moyen_mn"),
        "max": ("Tableaux 12-15 — Temps MAXIMAL observé", "temps_max_mn"),
    }
    for agregat, (titre_tab, champ) in libelles.items():
        ajouter_titre(titre_tab, niveau=2)
        # Group par tronçon
        par_tr: dict[str, dict[str, tuple[int | None, int]]] = {}
        for s in stats:
            nom = s.troncon_nom
            par_tr.setdefault(nom, {})
            par_tr[nom][s.type_jour] = (getattr(s, champ), s.nb_mesures)
        # Filtre : que les tronçons avec des mesures
        par_tr = {n: v for n, v in par_tr.items()
                  if sum(nb for _, nb in v.values()) > 0}
        if not par_tr:
            ajouter_paragraphe("Aucune mesure réelle sur cette campagne.", italique=True)
            doc.add_paragraph()
            continue
        tab = doc.add_table(rows=1, cols=4)
        tab.style = "Light Grid Accent 1"
        hdr = tab.rows[0].cells
        for i, txt in enumerate(("TRONÇON", "JOURS OUVRABLES (Min)",
                                  "WEEK-ENDS (Min)", "NB MESURES")):
            hdr[i].text = txt
            shade(hdr[i], "1A365D")
            for p in hdr[i].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = BLANC
                    run.font.size = Pt(9)
        for nom, vals in par_tr.items():
            jo_v, jo_n = vals.get("jour_ouvrable", (None, 0))
            we_v, we_n = vals.get("week_end", (None, 0))
            row = tab.add_row().cells
            row[0].text = nom
            row[1].text = "—" if jo_v is None else str(jo_v)
            row[2].text = "—" if we_v is None else str(we_v)
            row[3].text = str(jo_n + we_n)
            for c in row:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

    # ---- Graphiques 1-12 — BarChart par tronçon (min puis max) ----
    for agregat, titre_g in (("min", "Graphiques 1-6 — Temps MIN observé par jour"),
                              ("max", "Graphiques 7-12 — Temps MAX observé par jour")):
        ajouter_titre(titre_g, niveau=2)
        couleur = "#16a34a" if agregat == "min" else "#E74C3C"
        for t in troncons:
            serie = rapport_paa.serie_graphique(
                db, t.id, debut_utc, fin_utc, agregat=agregat,
                heure_debut=heure_debut, heure_fin=heure_fin,
            )
            if not serie:
                continue
            # Génération PNG matplotlib
            labels = [p.libelle_jour for p in serie]
            valeurs = [p.temps_mn for p in serie]
            fig, ax = plt.subplots(figsize=(8, 3), dpi=100)
            ax.bar(range(len(valeurs)), valeurs, color=couleur, edgecolor="white")
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=7)
            ax.set_ylabel("Temps (Min)", fontsize=8)
            ax.set_title(f"{t.nom} — Temps {agregat} (Min)", fontsize=9, color="#1A365D")
            ax.tick_params(axis="y", labelsize=7)
            ax.grid(axis="y", linestyle="--", alpha=0.3)
            fig.tight_layout()
            buf = io.BytesIO()
            fig.savefig(buf, format="png", bbox_inches="tight")
            plt.close(fig)
            buf.seek(0)
            doc.add_picture(buf, width=Cm(22))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # ---- Tableau 16 — Zones congestionnées ----
    ajouter_titre("Tableau 16 — Tronçons congestionnés (règles DEESP)", niveau=1)
    ajouter_paragraphe(
        f"Critère par mesure : couleur Google Maps — ROUGE OU ORANGE ≥ 50 %. "
        f"Seuils appliqués : ≥ {seuil_j} occurrence(s) / jour-indicatif OU "
        f"≥ {seuil_s} occurrence(s) / semaine.",
        italique=True,
    )
    if not cong:
        ajouter_paragraphe("Aucun tronçon congestionné sur cette campagne.",
                           italique=True)
    else:
        tab16 = doc.add_table(rows=1, cols=6)
        tab16.style = "Light Grid Accent 1"
        hdr = tab16.rows[0].cells
        for i, txt in enumerate(("AXE", "SOUS-TRONÇON", "TRANCHE",
                                  "NB/SEM.", "RÈGLE", "RÉPARTITION/JOUR")):
            hdr[i].text = txt
            shade(hdr[i], "1A365D")
            for p in hdr[i].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = BLANC
                    run.font.size = Pt(9)
        for c in cong:
            sous = (
                f"{c.sous_troncon_code} — {c.sous_troncon_nom or ''}"
                if c.sous_troncon_code
                else "axe entier"
            )
            regles = []
            if c.regle_jour_indicatif:
                regles.append(f"≥{seuil_j}/jour")
            if c.regle_semaine:
                regles.append(f"≥{seuil_s}/sem")
            row = tab16.add_row().cells
            row[0].text = c.troncon_nom or ""
            row[1].text = sous
            row[2].text = f"{c.heure:02d}h-{c.heure + 1:02d}h"
            row[3].text = str(c.nb_total_semaine)
            row[4].text = " | ".join(regles) or "—"
            row[5].text = " ".join(
                f"{j[:3]}:{n}" for j, n in c.nb_jours_congestionnes_par_type.items()
            )
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)

    # Pied de page
    doc.add_paragraph()
    pied = doc.add_paragraph()
    pr = pied.add_run(
        "Document généré en temps réel par FLUIDIS — "
        "Port Autonome d'Abidjan"
    )
    pr.italic = True
    pr.font.size = Pt(8)
    pr.font.color.rgb = GRIS
    pied.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sérialisation finale du document — étape la plus susceptible
    # d'échouer si une cellule contient une valeur non sérialisable.
    try:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        taille = len(buf.getvalue())
    except Exception:
        logger.exception("Echec doc.save() lors de la serialisation finale du .docx.")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Echec serialisation du document Word — voir logs serveur.",
        )

    nom = f"rapport_deesp_{campagne}.docx"
    logger.info(
        "GET /rapport/export/word OK — %s genere (%d octets, %d troncons).",
        nom, taille, len(troncons),
    )
    return Response(
        content=buf.getvalue(),
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{nom}"'},
    )


# ---------------------------------------------------------------------------
# GET /rapport/matrice-temps — matrice durées brutes heure × date
# ---------------------------------------------------------------------------


@router.get(
    "/matrice-temps",
    summary="Matrice de temps de traversée par créneau horaire × date",
    description=(
        "Pour le tronçon sélectionné et la plage de dates, renvoie pour chaque "
        "créneau horaire DEESP (07h-19h) et chaque date la durée de traversée "
        "observée (en secondes) — toutes sources confondues (google, terrain, "
        "historique). Permet de visualiser les temps réels sans agrégation."
    ),
)
async def get_matrice_temps(
    campagne: str = Query(..., description="Format 'AAAA-MM'."),
    troncon_id: int = Query(..., description="ID du tronçon à analyser."),
    debut: DateType | None = Query(None),
    fin: DateType | None = Query(None),
    heure_debut: int = Query(0, ge=0, le=23),
    heure_fin: int = Query(24, ge=1, le=24),
    sous_troncon_id: int | None = Query(None),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    from app.models.models import SousTroncon, Troncon, axe_sous_troncons as m2m
    logger.info(
        "GET /rapport/matrice-temps — troncon_id=%d sous=%s campagne=%r",
        troncon_id, sous_troncon_id, campagne,
    )
    troncon = db.get(Troncon, troncon_id)
    if troncon is None:
        raise HTTPException(status_code=404, detail=f"Tronçon {troncon_id} introuvable.")
    sous = None
    if sous_troncon_id is not None:
        sous = db.get(SousTroncon, sous_troncon_id)
        if sous is None:
            raise HTTPException(status_code=404, detail="Sous-tronçon introuvable.")
        # Vérifier parent principal OU rattachement M2M (multi-parent)
        if sous.troncon_id != troncon_id:
            from sqlalchemy import select as sa_select
            lien = db.execute(
                sa_select(m2m.c.axe_id).where(
                    m2m.c.axe_id == troncon_id,
                    m2m.c.sous_troncon_id == sous_troncon_id,
                )
            ).first()
            if lien is None:
                raise HTTPException(status_code=404, detail="Sous-tronçon introuvable.")
    debut_utc, fin_utc = _bornes_utc(campagne, debut, fin)
    result = rapport_paa.matrice_temps(
        db, troncon_id, debut_utc, fin_utc,
        heure_debut=heure_debut, heure_fin=heure_fin,
        sous_troncon_id=sous_troncon_id,
    )
    ref_dist = sous.distance_m if sous is not None else troncon.distance_ref_m
    ref_vit = 50.0 if sous is not None else troncon.vitesse_ref_kmh
    nom_affichage = (
        f"{troncon.nom} : {sous.nom_court} ({sous.code})"
        if sous is not None else troncon.nom
    )
    return {
        "troncon_nom": nom_affichage,
        "distance_m": ref_dist,
        "vitesse_ref_kmh": ref_vit,
        "temps_ref_s": round(ref_dist / (ref_vit * 1000 / 3600)),
        **result,
    }


# ---------------------------------------------------------------------------
# POST /rapport/import-mesures-excel — import Excel/CSV de mesures manuelles
# ---------------------------------------------------------------------------


@router.post(
    "/import-mesures-excel",
    summary="Import Excel/CSV de mesures de temps de traversée",
    description=(
        "Accepte un fichier Excel (.xlsx, .xls) ou CSV (.csv) avec les colonnes :\n\n"
        "- `date` (YYYY-MM-DD ou DD/MM/YYYY)\n"
        "- `heure` (entier 0-23)\n"
        "- `troncon_id` (entier)\n"
        "- `duree_mn` (décimal, minutes)\n\n"
        "Insère dans la table `mesures` avec `source=historique_paa_2025`. "
        "Les doublons (même troncon + même horodatage exact) sont ignorés silencieusement."
    ),
)
async def import_mesures_excel(
    fichier: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    import io
    from datetime import time as dtime
    from zoneinfo import ZoneInfo

    import pandas as pd
    from sqlalchemy import select as sa_select

    from app.models.models import Mesure, SourceMesure, Troncon

    logger.info("POST /rapport/import-mesures-excel — fichier=%s", fichier.filename)

    contenu = await fichier.read()
    nom = (fichier.filename or "").lower()
    try:
        if nom.endswith(".csv"):
            df = pd.read_csv(io.BytesIO(contenu))
        else:
            df = pd.read_excel(io.BytesIO(contenu))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Lecture du fichier échouée : {exc}")

    # Normalisation des noms de colonnes (minuscules, espaces → underscore)
    df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]
    cols_requises = {"date", "heure", "troncon_id", "duree_mn"}
    manquantes = cols_requises - set(df.columns)
    if manquantes:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Colonnes manquantes : {', '.join(sorted(manquantes))}. "
                f"Colonnes trouvées : {', '.join(df.columns.tolist())}."
            ),
        )

    fuseau = ZoneInfo(get_settings().tz)
    nb_inserees = 0
    nb_doublons = 0
    erreurs: list[str] = []

    for i, row in df.iterrows():
        try:
            date_val = pd.to_datetime(row["date"]).date()
            heure = int(row["heure"])
            if not (0 <= heure <= 23):
                erreurs.append(f"Ligne {i + 2} : heure {heure} invalide (0-23)")
                continue
            troncon_id_val = int(row["troncon_id"])
            duree_mn = float(row["duree_mn"])
            if duree_mn <= 0:
                erreurs.append(f"Ligne {i + 2} : duree_mn invalide ({duree_mn})")
                continue
            duree_s = int(round(duree_mn * 60))

            dt_local = datetime.combine(date_val, dtime(heure, 0, 0), tzinfo=fuseau)
            dt_utc = dt_local.astimezone(timezone.utc)

            troncon = db.get(Troncon, troncon_id_val)
            if troncon is None:
                erreurs.append(f"Ligne {i + 2} : tronçon {troncon_id_val} introuvable")
                continue

            # Doublon : même tronçon + même horodatage exact
            existe = db.execute(
                sa_select(Mesure.id).where(
                    Mesure.troncon_id == troncon_id_val,
                    Mesure.source == SourceMesure.historique_paa_2025,
                    Mesure.horodatage == dt_utc,
                ).limit(1)
            ).first()
            if existe:
                nb_doublons += 1
                continue

            vitesse_kmh = round(troncon.distance_m / duree_s * 3.6, 2) if duree_s > 0 else 0.0
            db.add(Mesure(
                troncon_id=troncon_id_val,
                horodatage=dt_utc,
                duree_trafic_s=duree_s,
                duree_sans_trafic_s=None,
                source=SourceMesure.historique_paa_2025,
                vitesse_moyenne_kmh=vitesse_kmh,
                aberrante=False,
                est_congestionne=None,
                pourcentage_rouge=None,
                pourcentage_orange=None,
                pourcentage_vert=None,
            ))
            nb_inserees += 1
        except Exception as exc:
            erreurs.append(f"Ligne {i + 2} : {exc}")

    if nb_inserees > 0:
        db.commit()

    logger.info(
        "import-mesures-excel : %d insérées, %d doublons, %d erreurs",
        nb_inserees, nb_doublons, len(erreurs),
    )
    return {
        "nb_inserees": nb_inserees,
        "nb_doublons": nb_doublons,
        "erreurs": erreurs[:20],
        "message": f"{nb_inserees} mesure(s) importée(s), {nb_doublons} doublon(s) ignoré(s).",
    }


# ---------------------------------------------------------------------------
# GET /rapport/graphique/{troncon_id}
# ---------------------------------------------------------------------------


@router.get(
    "/graphique/{troncon_id}",
    summary="Données BarChart pour Graphiques 1-12 (min ou max par jour)",
    description=(
        "Renvoie une série prête à tracer en BarChart Recharts. Chaque "
        "point = un jour de la campagne, hauteur = temps min ou max observé "
        "ce jour-là sur ce tronçon. C'est le format attendu pour les "
        "Graphiques 1, 3, 5 (temps min sens aller), 2, 4, 6 (temps min "
        "sens retour), 7, 9, 11 (temps max aller) et 8, 10, 12 (temps max "
        "retour) du rapport."
    ),
)
async def get_graphique(
    troncon_id: int,
    campagne: str = Query(..., description="Format 'AAAA-MM'."),
    agregat: str = Query("min", description="`min` ou `max`."),
    heure_debut: int = Query(0, ge=0, le=23),
    heure_fin: int = Query(24, ge=1, le=24),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    if agregat not in ("min", "max"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="`agregat` doit valoir 'min' ou 'max'.",
        )
    debut_utc, fin_utc = _parser_campagne(campagne)
    serie = rapport_paa.serie_graphique(
        db, troncon_id, debut_utc, fin_utc, agregat=agregat,
        heure_debut=heure_debut, heure_fin=heure_fin,
    )
    return {
        "troncon_id": troncon_id,
        "campagne": campagne,
        "agregat": agregat,
        "axe_y_unite": "minutes",
        "nb_points": len(serie),
        "points": [
            {
                "date": p.date_locale,
                "libelle_jour": p.libelle_jour,
                "temps_mn": p.temps_mn,
            }
            for p in serie
        ],
    }


# ---------------------------------------------------------------------------
# GET /rapport/comparaison
# ---------------------------------------------------------------------------


@router.get(
    "/comparaison",
    summary="Tableau 19 — Comparaison pluriannuelle entre 2 campagnes",
    description=(
        "Reproduit le Tableau 19 du rapport (comparaison fév 2025 vs "
        "oct 2025). Pour chaque tronçon × type-jour, renvoie min/moyen/max "
        "des deux campagnes et le delta du temps moyen."
    ),
)
async def get_comparaison(
    campagne_a: str = Query(
        ..., description="Campagne de référence, format 'AAAA-MM'.",
    ),
    campagne_b: str = Query(
        ..., description="Campagne comparée, format 'AAAA-MM'.",
    ),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    debut_a_utc, fin_a_utc = _parser_campagne(campagne_a)
    debut_b_utc, fin_b_utc = _parser_campagne(campagne_b)
    # rapport_paa.comparaison_campagnes attend des dates locales — on
    # reconstruit depuis nos UTC.
    fuseau_local = ZoneInfo(get_settings().tz)
    debut_a = debut_a_utc.astimezone(fuseau_local).date()
    fin_a = fin_a_utc.astimezone(fuseau_local).date()
    debut_b = debut_b_utc.astimezone(fuseau_local).date()
    fin_b = fin_b_utc.astimezone(fuseau_local).date()
    lignes = rapport_paa.comparaison_campagnes(
        db, (debut_a, fin_a), (debut_b, fin_b),
    )
    return {
        "campagne_a": campagne_a,
        "campagne_b": campagne_b,
        "nb_lignes": len(lignes),
        "lignes": lignes,
    }
