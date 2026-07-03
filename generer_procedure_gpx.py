"""Génère le document Word de procédure terrain GPX pour FLUIDIS.

Lancement :
    py -3.11 generer_procedure_gpx.py
ou bien :
    "C:\\Users\\hp\\AppData\\Local\\Programs\\Python\\Python311\\python.exe" generer_procedure_gpx.py

Sortie : Procedure_Collecte_GPX_Terrain_FLUIDIS.docx
"""

from datetime import date
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Couleurs et constantes de mise en forme
# ---------------------------------------------------------------------------
COULEUR_BLEU = RGBColor(0x1F, 0x4E, 0x79)
COULEUR_BLEU_CLAIR = "D9E2F3"
COULEUR_GRIS_CLAIR = "F2F2F2"
COULEUR_GRIS = RGBColor(0x59, 0x59, 0x59)
COULEUR_BLANC = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Helpers de mise en forme bas niveau
# ---------------------------------------------------------------------------
def shade_cell(cell, hex_color):
    """Applique une couleur de fond à une cellule (XML direct car python-docx ne l'expose pas)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BFBFBF", size="4"):
    """Pose des bordures fines grises sur les 4 côtés d'une cellule."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for direction in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{direction}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def ajouter_paragraphe(doc, texte, gras=False, italique=False, taille=11,
                      couleur=None, alignement=None, espace_apres=4):
    p = doc.add_paragraph()
    if alignement is not None:
        p.alignment = alignement
    p.paragraph_format.space_after = Pt(espace_apres)
    run = p.add_run(texte)
    run.font.size = Pt(taille)
    run.bold = gras
    run.italic = italique
    if couleur is not None:
        run.font.color.rgb = couleur
    return p


def ajouter_h1(doc, texte):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(texte)
    run.font.color.rgb = COULEUR_BLEU
    run.font.size = Pt(18)
    run.bold = True
    return p


def ajouter_h2(doc, texte):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(texte)
    run.font.color.rgb = COULEUR_BLEU
    run.font.size = Pt(14)
    run.bold = True
    return p


def ajouter_h3(doc, texte):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    run = p.add_run(texte)
    run.font.color.rgb = COULEUR_GRIS
    run.font.size = Pt(12)
    run.bold = True
    return p


def ajouter_puce(doc, texte, gras_segments=None):
    """Ajoute un paragraphe à puce. gras_segments = liste de (texte, gras)."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if gras_segments is None:
        p.add_run(texte)
    else:
        for segment, gras in gras_segments:
            run = p.add_run(segment)
            run.bold = gras
    return p


def ajouter_numerote(doc, texte, gras_segments=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    if gras_segments is None:
        p.add_run(texte)
    else:
        for segment, gras in gras_segments:
            run = p.add_run(segment)
            run.bold = gras
    return p


def encadre_info(doc, titre, lignes, couleur_fond=COULEUR_BLEU_CLAIR):
    """Encadre informatif : 1 cellule avec un titre en gras puis des lignes."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    shade_cell(cell, couleur_fond)
    set_cell_borders(cell, color="A6BBD9", size="6")
    # Premier paragraphe = titre
    cell.paragraphs[0].text = ""
    run_titre = cell.paragraphs[0].add_run(titre)
    run_titre.bold = True
    run_titre.font.color.rgb = COULEUR_BLEU
    run_titre.font.size = Pt(11)
    # Lignes suivantes
    for ligne in lignes:
        p = cell.add_paragraph(ligne)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(11)
    # Espace après l'encadré
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def ajouter_tableau(doc, entetes, lignes, largeurs_cm=None, premiere_col_grise=False):
    """Crée un tableau avec un en-tête bleu et des lignes alternées éventuelles."""
    table = doc.add_table(rows=1 + len(lignes), cols=len(entetes))
    table.autofit = False
    if largeurs_cm:
        for ligne in table.rows:
            for cell, larg in zip(ligne.cells, largeurs_cm):
                cell.width = Cm(larg)

    # En-têtes
    for cell, entete in zip(table.rows[0].cells, entetes):
        cell.text = ""
        shade_cell(cell, "1F4E79")
        set_cell_borders(cell)
        run = cell.paragraphs[0].add_run(entete)
        run.bold = True
        run.font.color.rgb = COULEUR_BLANC
        run.font.size = Pt(11)

    # Lignes
    for idx_ligne, valeurs in enumerate(lignes):
        for idx_cell, val in enumerate(valeurs):
            cell = table.rows[idx_ligne + 1].cells[idx_cell]
            cell.text = ""
            if premiere_col_grise and idx_cell == 0:
                shade_cell(cell, COULEUR_GRIS_CLAIR)
            set_cell_borders(cell)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(11)
            if premiere_col_grise and idx_cell == 0:
                run.bold = True
    return table


def saut_de_page(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Données : les 6 tronçons
# ---------------------------------------------------------------------------
TRONCONS = [
    {
        "num": "1A",
        "nom": "CARENA (Plateau) → Pharmacie Vridi Palm Beach",
        "distance": "14,9 km",
        "origine": "Carena (chantier naval) — Plateau, Boulevard de Marseille",
        "lat_o": "5.328119", "lon_o": "-4.028563",
        "destination": "Pharmacie Vridi Palm Beach — Marcory, Boulevard Hortense Aka Anghui",
        "lat_d": "5.258705", "lon_d": "-3.981960",
        "duree": "≈ 18 min (sans bouchon, à 50 km/h)",
    },
    {
        "num": "1B",
        "nom": "Pharmacie Vridi Palm Beach → CARENA (Plateau)",
        "distance": "14,9 km",
        "origine": "Pharmacie Vridi Palm Beach — Marcory, Boulevard Hortense Aka Anghui",
        "lat_o": "5.258705", "lon_o": "-3.981960",
        "destination": "Carena (chantier naval) — Plateau, Boulevard de Marseille",
        "lat_d": "5.328119", "lon_d": "-4.028563",
        "duree": "≈ 18 min (sans bouchon, à 50 km/h)",
    },
    {
        "num": "2A",
        "nom": "Toyota CFAO (Treichville) → Pharmacie Vridi Palm Beach",
        "distance": "8,0 km",
        "origine": "Toyota CFAO — Treichville, Rue Pasteur",
        "lat_o": "5.295971", "lon_o": "-4.005131",
        "destination": "Pharmacie Vridi Palm Beach — Marcory, Boulevard Hortense Aka Anghui",
        "lat_d": "5.258705", "lon_d": "-3.981960",
        "duree": "≈ 10 min (sans bouchon, à 50 km/h)",
    },
    {
        "num": "2B",
        "nom": "Pharmacie Vridi Palm Beach → Toyota CFAO (Treichville)",
        "distance": "8,0 km",
        "origine": "Pharmacie Vridi Palm Beach — Marcory, Boulevard Hortense Aka Anghui",
        "lat_o": "5.258705", "lon_o": "-3.981960",
        "destination": "Toyota CFAO — Treichville, Rue Pasteur",
        "lat_d": "5.295971", "lon_d": "-4.005131",
        "duree": "≈ 10 min (sans bouchon, à 50 km/h)",
    },
    {
        "num": "3A",
        "nom": "Agence SODECI Zone 4 → Pharmacie Vridi Palm Beach",
        "distance": "8,3 km",
        "origine": "Agence SODECI Zone 4 — Treichville/Marcory",
        "lat_o": "5.293686", "lon_o": "-4.000390",
        "destination": "Pharmacie Vridi Palm Beach — Marcory, Boulevard Hortense Aka Anghui",
        "lat_d": "5.258705", "lon_d": "-3.981960",
        "duree": "≈ 10 min (sans bouchon, à 50 km/h)",
    },
    {
        "num": "3B",
        "nom": "Pharmacie Vridi Palm Beach → Agence SODECI Zone 4",
        "distance": "8,3 km",
        "origine": "Pharmacie Vridi Palm Beach — Marcory, Boulevard Hortense Aka Anghui",
        "lat_o": "5.258705", "lon_o": "-3.981960",
        "destination": "Agence SODECI Zone 4 — Treichville/Marcory",
        "lat_d": "5.293686", "lon_d": "-4.000390",
        "duree": "≈ 10 min (sans bouchon, à 50 km/h)",
    },
]


def fiche_troncon(doc, t):
    """Fiche détaillée d'un tronçon : encadré + tableau coordonnées + zones à remplir."""
    ajouter_h2(doc, f"Tronçon {t['num']} — {t['nom']}")

    ajouter_tableau(
        doc,
        entetes=["Champ", "Valeur"],
        lignes=[
            ("Distance officielle", t["distance"]),
            ("Point de départ", t["origine"]),
            ("Coordonnées départ", f"{t['lat_o']}, {t['lon_o']}"),
            ("Point d'arrivée", t["destination"]),
            ("Coordonnées arrivée", f"{t['lat_d']}, {t['lon_d']}"),
            ("Durée idéale", t["duree"]),
        ],
        largeurs_cm=[4.5, 11.5],
        premiere_col_grise=True,
    )

    doc.add_paragraph()
    ajouter_paragraphe(doc, "À remplir sur place :", gras=True, taille=11)

    ajouter_tableau(
        doc,
        entetes=["Heure de départ", "Heure d'arrivée", "Durée mesurée"],
        lignes=[("__h__min__s", "__h__min__s", "____min____s")],
        largeurs_cm=[5.3, 5.3, 5.4],
    )

    doc.add_paragraph()
    ajouter_paragraphe(doc, "Nom du fichier GPX à enregistrer :", gras=True, taille=11)
    p = doc.add_paragraph()
    run = p.add_run(f"paa_t{t['num']}_AAAAMMJJ_HHMM.gpx")
    run.font.name = "Consolas"
    run.font.color.rgb = COULEUR_BLEU
    run.font.size = Pt(11)
    run.bold = True

    ajouter_paragraphe(doc, "Observations terrain (météo, incidents, fluidité) :", gras=True, taille=11)
    for _ in range(3):
        ajouter_paragraphe(doc, "_" * 75, taille=11, espace_apres=2)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Construction du document
# ---------------------------------------------------------------------------
def generer():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Police par défaut
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)

    # En-tête et pied de page
    header = doc.sections[0].header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("FLUIDIS — Procédure GPX terrain")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Port Autonome d'Abidjan — Phase P5 (validation terrain)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ----- Page de garde -----
    for _ in range(4):
        doc.add_paragraph()
    ajouter_paragraphe(doc, "FLUIDIS", gras=True, taille=36,
                       couleur=COULEUR_BLEU,
                       alignement=WD_ALIGN_PARAGRAPH.CENTER, espace_apres=12)
    ajouter_paragraphe(doc, "Port Autonome d'Abidjan", taille=18,
                       couleur=COULEUR_GRIS,
                       alignement=WD_ALIGN_PARAGRAPH.CENTER, espace_apres=36)
    ajouter_paragraphe(doc, "Procédure de collecte terrain GPX",
                       gras=True, taille=24,
                       alignement=WD_ALIGN_PARAGRAPH.CENTER, espace_apres=12)
    ajouter_paragraphe(doc, "Phase P5 — validation hebdomadaire des sources API",
                       italique=True, taille=12, couleur=COULEUR_GRIS,
                       alignement=WD_ALIGN_PARAGRAPH.CENTER, espace_apres=36)

    aujourdhui = date.today()
    mois_fr = ["janvier", "février", "mars", "avril", "mai", "juin",
               "juillet", "août", "septembre", "octobre", "novembre", "décembre"]
    date_affichage = f"{aujourdhui.day} {mois_fr[aujourdhui.month - 1]} {aujourdhui.year}"

    ajouter_paragraphe(doc, f"Sortie terrain prévue : {date_affichage}",
                       taille=12,
                       alignement=WD_ALIGN_PARAGRAPH.CENTER, espace_apres=4)
    ajouter_paragraphe(doc, "6 tronçons dirigés — 3 axes aller-retour",
                       taille=12, couleur=COULEUR_GRIS,
                       alignement=WD_ALIGN_PARAGRAPH.CENTER, espace_apres=24)

    for _ in range(3):
        doc.add_paragraph()

    encadre_info(doc, "À lire avant de partir", [
        "Cette procédure décrit le protocole exact à suivre pour enregistrer une trace GPX réelle sur chacun des 6 tronçons surveillés par le projet FLUIDIS.",
        "Les fichiers GPX produits serviront à confronter les durées annoncées par Google Routes API à la réalité terrain (Phase P5 du projet).",
        "Lire la procédure en entier avant de partir. Compter une demi-journée pour l'ensemble des 6 trajets.",
    ])

    saut_de_page(doc)

    # ----- 1. Objectif -----
    ajouter_h1(doc, "1. Objectif de la sortie terrain")
    ajouter_paragraphe(doc,
        "L'application FLUIDIS collecte automatiquement, toutes les 20 minutes "
        "entre 7h et 19h, les temps de parcours annoncés par Google Routes API sur les "
        "6 tronçons du port. Pour garantir la fiabilité de cette source, le cahier des "
        "charges impose une confrontation hebdomadaire avec une mesure terrain réelle."
    )
    ajouter_paragraphe(doc, "Le but de cette sortie est donc simple :")
    ajouter_puce(doc, "", [("parcourir physiquement ", False), ("chacun des 6 tronçons", True)])
    ajouter_puce(doc, "", [("enregistrer la ", False), ("trace GPS complète (fichier GPX)", True), (" pendant le trajet", False)])
    ajouter_puce(doc, "", [("noter ", False), ("l'heure exacte de départ et d'arrivée", True)])
    ajouter_puce(doc, "transmettre les fichiers au backend (import en P5)")

    ajouter_h2(doc, "Pourquoi un GPX et pas seulement une heure ?")
    ajouter_paragraphe(doc,
        "Une simple heure de départ et d'arrivée pourrait suffire pour calculer une durée — "
        "mais le GPX apporte deux garanties supplémentaires :"
    )
    ajouter_numerote(doc, "Il prouve que l'itinéraire emprunté correspond bien à l'axe officiel (pas de raccourci ni de détour).")
    ajouter_numerote(doc, "Il permet de détecter d'éventuelles pertes de signal GPS (tunnels, immeubles), ce qui invalide la mesure.")
    ajouter_numerote(doc, "Il pourra être réutilisé plus tard pour analyser à quel endroit précis du tronçon le trafic ralentit.")

    saut_de_page(doc)

    # ----- 2. Matériel -----
    ajouter_h1(doc, "2. Matériel nécessaire")
    for item in [
        "Smartphone Android avec GPS activé (Android 7 ou supérieur)",
        "Batterie chargée à 100 % au départ (ou batterie externe)",
        "Connexion Internet mobile activée (recommandé, pas obligatoire pour le GPS)",
        "Véhicule (voiture, taxi ou véhicule de service)",
        "Un assistant pour noter les heures pendant que le conducteur conduit (optionnel mais recommandé)",
        "Cahier ou ce document imprimé pour noter les observations terrain",
    ]:
        ajouter_puce(doc, item)

    # ----- 3. Préparation -----
    ajouter_h1(doc, "3. Préparation préalable")

    ajouter_h2(doc, "3.1 Installer l'application GPS Logger")
    ajouter_paragraphe(doc,
        "L'application recommandée est GPS Logger (gratuite, open-source, sans publicité, "
        "faible consommation batterie). Elle exporte directement au format GPX standard."
    )
    ajouter_puce(doc, "Ouvrir le Google Play Store sur l'Android")
    ajouter_puce(doc, "", [("Chercher : ", False), ("GPSLogger for Android", True), (" (auteur Mendhak — icône triangle bleu)", False)])
    ajouter_puce(doc, "Installer et ouvrir l'application")
    ajouter_puce(doc, "Accepter les permissions de localisation (« Toujours autoriser » pour fonctionner en arrière-plan)")

    ajouter_paragraphe(doc, "Alternatives possibles si GPSLogger n'est pas disponible :", gras=True)
    ajouter_puce(doc, "OsmAnd (gratuit, plus complet, sait aussi exporter en GPX)")
    ajouter_puce(doc, "Geo Tracker (gratuit, interface très simple)")
    ajouter_puce(doc, "Strava (nécessite un compte, exporter le GPX depuis le site web après la sortie)")

    ajouter_h2(doc, "3.2 Configurer GPS Logger pour la sortie")
    ajouter_paragraphe(doc,
        "Dans GPS Logger, ouvrir le menu (icône en haut à gauche) puis aller dans « Préférences de "
        "journalisation » et appliquer ces réglages :"
    )

    ajouter_tableau(
        doc,
        entetes=["Paramètre", "Valeur à régler"],
        lignes=[
            ("Intervalle de journalisation", "1 seconde (le plus précis)"),
            ("Distance minimale entre points", "0 mètre (enregistrer même à l'arrêt)"),
            ("Précision GPS exigée", "40 mètres (compromis ville/précision)"),
            ("Démarrage automatique au boot", "Désactivé"),
            ("Garder l'écran allumé pendant l'enregistrement", "Activé (évite les coupures)"),
            ("Format d'enregistrement", "GPX (uniquement) — désactiver KML et autres"),
            ("Création d'un nouveau fichier", "À chaque démarrage de journalisation"),
        ],
        largeurs_cm=[8.0, 8.0],
        premiere_col_grise=True,
    )

    doc.add_paragraph()
    ajouter_h2(doc, "3.3 Vérifier la qualité du GPS avant de démarrer")
    ajouter_paragraphe(doc, "Sortir à l'air libre (ne pas démarrer depuis un parking couvert).")
    ajouter_numerote(doc, "Ouvrir GPS Logger.")
    ajouter_numerote(doc, "Attendre que la précision affichée descende à 10 mètres ou moins (peut prendre 30 secondes).")
    ajouter_numerote(doc, "", [
        ("Si la précision reste supérieure à 50 mètres : ", False),
        ("ne pas démarrer", True),
        (", attendre encore ou se déplacer vers un endroit plus dégagé.", False),
    ])

    saut_de_page(doc)

    # ----- 4. Procédure générale -----
    ajouter_h1(doc, "4. Procédure générale à appliquer pour chaque tronçon")
    ajouter_paragraphe(doc,
        "La même procédure est appliquée pour les 6 tronçons. Compter environ 30 à 45 minutes "
        "par tronçon (trajet + pause), soit environ 4 à 5 heures pour l'ensemble."
    )

    ajouter_h2(doc, "Étapes à suivre pour chaque trajet")
    etapes = [
        [("Se positionner au ", False), ("point de départ exact", True), (" (vérifier avec les coordonnées GPS de la fiche du tronçon).", False)],
        [("Vérifier que la ", False), ("précision GPS", True), (" affichée est inférieure ou égale à 10 mètres.", False)],
        [("Dans GPS Logger : appuyer sur le bouton ", False), ("« Démarrer la journalisation »", True), (" (bouton vert avec icône lecture).", False)],
        [("Noter immédiatement ", False), ("l'heure exacte de démarrage", True), (" sur la fiche du tronçon (à la seconde près si possible).", False)],
        [("Conduire ", False), ("normalement", True), (", en respectant le code de la route. Suivre l'itinéraire habituel des camions et véhicules empruntant cet axe.", False)],
        [("À l'arrivée au ", False), ("point de destination exact", True), (", appuyer immédiatement sur ", False), ("« Arrêter la journalisation »", True), (".", False)],
        [("Noter ", False), ("l'heure exacte d'arrivée", True), (" sur la fiche.", False)],
        [("Renommer immédiatement le fichier GPX selon la convention paa_tXY_AAAAMMJJ_HHMM.gpx (voir section 5).", False)],
        [("Inscrire les ", False), ("observations terrain", True), (" (fluidité, météo, incidents constatés).", False)],
        [("Marquer une pause de 5 à 10 minutes avant de démarrer le tronçon suivant.", False)],
    ]
    for segments in etapes:
        ajouter_numerote(doc, "", segments)

    ajouter_h2(doc, "Règles d'or à respecter")
    encadre_info(doc, "Cinq règles à ne jamais oublier", [
        "1. Ne JAMAIS prendre un raccourci ou un itinéraire différent : seul l'axe officiel est valable.",
        "2. Démarrer et arrêter la journalisation AUX POINTS EXACTS — pas 200 m avant ou après.",
        "3. Faire les trajets entre 7h00 et 19h00 (Africa/Abidjan) pour pouvoir comparer aux mesures Google.",
        "4. Ne PAS arrêter et redémarrer la journalisation en cours de trajet (sinon le GPX est coupé en deux).",
        "5. En cas de problème majeur (accident bloquant la route, déviation imposée), noter sur la fiche et reprendre le trajet plus tard ou un autre jour.",
    ])

    saut_de_page(doc)

    # ----- 5. Nommage -----
    ajouter_h1(doc, "5. Convention de nommage des fichiers GPX")
    ajouter_paragraphe(doc,
        "Pour faciliter l'import en base, chaque fichier GPX doit être renommé selon une "
        "convention stricte :"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("paa_tXY_AAAAMMJJ_HHMM.gpx")
    run.font.name = "Consolas"
    run.font.size = Pt(16)
    run.font.color.rgb = COULEUR_BLEU
    run.bold = True

    ajouter_tableau(
        doc,
        entetes=["Élément", "Signification"],
        lignes=[
            ("paa_t", "Préfixe fixe pour identifier les fichiers FLUIDIS"),
            ("XY", "Numéro du tronçon : 1A, 1B, 2A, 2B, 3A, 3B"),
            ("AAAAMMJJ", "Date au format année-mois-jour (ex. 20260619 pour le 19 juin 2026)"),
            ("HHMM", "Heure de départ du trajet (ex. 0814 pour 8h14)"),
            (".gpx", "Extension du fichier (obligatoire)"),
        ],
        largeurs_cm=[4.0, 12.0],
        premiere_col_grise=True,
    )

    doc.add_paragraph()
    ajouter_paragraphe(doc, "Exemples concrets :", gras=True)
    for nom, description in [
        ("paa_t1A_20260619_0814.gpx", "  → Tronçon 1A (CARENA→Palm Beach), 19 juin 2026, départ 8h14"),
        ("paa_t2B_20260619_0945.gpx", "  → Tronçon 2B (Palm Beach→Toyota), 19 juin 2026, départ 9h45"),
        ("paa_t3A_20260619_1530.gpx", "  → Tronçon 3A (SODECI→Palm Beach), 19 juin 2026, départ 15h30"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(nom)
        run.font.name = "Consolas"
        run.font.color.rgb = COULEUR_BLEU
        p.add_run(description)

    saut_de_page(doc)

    # ----- 6. Fiches détaillées -----
    ajouter_h1(doc, "6. Fiches détaillées des 6 tronçons")
    ajouter_paragraphe(doc,
        "Une fiche par tronçon, à remplir au fur et à mesure de la sortie. Compléter "
        "les heures de départ et d'arrivée DÈS l'arrêt de la journalisation."
    )
    doc.add_paragraph()

    for t in TRONCONS:
        fiche_troncon(doc, t)

    saut_de_page(doc)

    # ----- 7. Vérifications -----
    ajouter_h1(doc, "7. Vérifications après la sortie")
    ajouter_paragraphe(doc, "De retour au bureau ou à la maison, avant de transmettre les fichiers :")

    ajouter_h2(doc, "7.1 Récupérer les fichiers depuis le smartphone")
    ajouter_paragraphe(doc, "Les fichiers GPX sont enregistrés par défaut dans le dossier suivant sur le smartphone :")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("/storage/emulated/0/GPSLogger/")
    run.font.name = "Consolas"
    run.font.color.rgb = COULEUR_BLEU
    run.font.size = Pt(11)

    ajouter_paragraphe(doc, "Trois méthodes pour les récupérer :")
    ajouter_puce(doc, "", [
        ("Méthode 1 — câble USB : ", True),
        ("brancher le smartphone à un ordinateur, ouvrir le stockage, copier les fichiers GPX dans un dossier du PC.", False),
    ])
    ajouter_puce(doc, "", [
        ("Méthode 2 — email : ", True),
        ("depuis GPS Logger, ouvrir chaque fichier et utiliser l'option « Partager → Email ».", False),
    ])
    ajouter_puce(doc, "", [
        ("Méthode 3 — Google Drive : ", True),
        ("synchroniser le dossier GPSLogger avec Google Drive, puis télécharger sur le PC.", False),
    ])

    ajouter_h2(doc, "7.2 Vérifier chaque fichier GPX")
    ajouter_paragraphe(doc,
        "Ouvrir chaque fichier avec un visualiseur (par exemple gpx.studio en ligne, ou OsmAnd "
        "sur Android) et vérifier :"
    )
    for item in [
        "Le tracé suit bien l'axe officiel du tronçon (pas de raccourci visible).",
        "Il n'y a pas de coupure (ligne droite suspecte entre deux points éloignés = perte GPS).",
        "La durée totale est cohérente avec ce qu'on a vécu (ex. 18 min, pas 2 min ni 3 heures).",
        "Le point de départ et le point d'arrivée sont bien situés aux lieux officiels (à 100 m près maximum).",
    ]:
        ajouter_puce(doc, item)

    ajouter_h2(doc, "7.3 Que faire si un GPX est défectueux ?")
    ajouter_puce(doc, "", [
        ("Coupure GPS importante (> 30 secondes) : ", True),
        ("noter sur la fiche du tronçon que la mesure est invalide et programmer une nouvelle sortie sur ce tronçon.", False),
    ])
    ajouter_puce(doc, "", [
        ("Itinéraire visiblement différent : ", True),
        ("même chose, mesure à refaire.", False),
    ])
    ajouter_puce(doc, "", [
        ("Coupure courte (< 10 secondes) : ", True),
        ("la mesure reste exploitable, garder le fichier.", False),
    ])

    # ----- 8. Transmission -----
    ajouter_h1(doc, "8. Transmission des fichiers")
    ajouter_paragraphe(doc, "Une fois les fichiers vérifiés, les transmettre selon l'une de ces modalités :")
    ajouter_numerote(doc, "Les envoyer par email au responsable du projet FLUIDIS.")
    ajouter_numerote(doc, "Les déposer dans un dossier partagé Google Drive ou OneDrive convenu.")
    ajouter_numerote(doc, "Quand l'endpoint d'import sera prêt (P5) : les uploader directement via la documentation Swagger du backend.")

    ajouter_paragraphe(doc, "À joindre obligatoirement à la livraison :", gras=True)
    ajouter_puce(doc, "Les 6 fichiers GPX (un par tronçon parcouru)")
    ajouter_puce(doc, "Cette procédure complétée (heures de départ/arrivée + observations)")
    ajouter_puce(doc, "Un court résumé : conditions météo de la journée, événements particuliers (manifestation, accident, etc.)")

    saut_de_page(doc)

    # ----- 9. Bonnes pratiques -----
    ajouter_h1(doc, "9. Bonnes pratiques et conseils")

    ajouter_h2(doc, "9.1 Choisir un bon créneau horaire")
    ajouter_paragraphe(doc,
        "Pour que la confrontation soit pertinente, faire les trajets pendant que le robot "
        "collecte (7h-19h Africa/Abidjan). Privilégier :"
    )
    ajouter_puce(doc, "Un jour ouvrable (mardi, mercredi ou jeudi, hors jour férié)")
    ajouter_puce(doc, "Le matin (8h-10h) pour observer le trafic d'entrée vers le port")
    ajouter_puce(doc, "Ou l'après-midi (16h-18h) pour observer le trafic de sortie")
    ajouter_puce(doc, "Éviter les heures de fermeture des écoles (12h, 12h30, 16h30) qui créent des pics atypiques")

    ajouter_h2(doc, "9.2 Ordre des trajets recommandé")
    ajouter_paragraphe(doc,
        "Pour limiter les déplacements à vide entre deux trajets, suivre cet ordre (en boucle "
        "autour du port) :"
    )
    ajouter_tableau(
        doc,
        entetes=["Ordre", "Tronçon", "Astuce"],
        lignes=[
            ("1", "1A — CARENA → Palm Beach", "Démarrer depuis le Plateau"),
            ("2", "3B — Palm Beach → SODECI Zone 4", "Repartir directement de Palm Beach"),
            ("3", "3A — SODECI → Palm Beach", "Trajet inverse depuis SODECI"),
            ("4", "2B — Palm Beach → Toyota CFAO", "Direction Treichville"),
            ("5", "2A — Toyota CFAO → Palm Beach", "Retour"),
            ("6", "1B — Palm Beach → CARENA", "Dernier trajet, retour vers le Plateau"),
        ],
        largeurs_cm=[1.5, 6.0, 8.5],
        premiere_col_grise=True,
    )

    doc.add_paragraph()
    ajouter_h2(doc, "9.3 Pendant le trajet")
    for item in [
        "Garder le smartphone dans une position stable (support voiture ou poche poitrine), pas posé dans la portière.",
        "Ne pas mettre l'appareil en mode économie d'énergie agressif (il couperait le GPS).",
        "Vérifier de temps en temps que la journalisation est toujours active (notification persistante).",
        "Respecter les limitations de vitesse — l'objectif est de mesurer un trajet « normal », pas record.",
    ]:
        ajouter_puce(doc, item)

    ajouter_h2(doc, "9.4 Gestion de la batterie")
    for item in [
        "Le GPS continu consomme 10-15 % de batterie par heure. Pour 6 trajets sur la journée, prévoir une batterie externe.",
        "Désactiver les applications en arrière-plan pour économiser.",
        "Si possible, recharger entre deux trajets.",
    ]:
        ajouter_puce(doc, item)

    # ----- 10. Problèmes -----
    ajouter_h1(doc, "10. Que faire en cas de problème")
    ajouter_tableau(
        doc,
        entetes=["Problème rencontré", "Action à prendre"],
        lignes=[
            ("Précision GPS très mauvaise (> 50 m) malgré l'attente",
             "Se déplacer vers une zone plus dégagée (sortir d'un canyon urbain, s'écarter des grands bâtiments)."),
            ("GPS Logger plante en cours de trajet",
             "Redémarrer l'app, refaire le trajet depuis le point de départ. Ne pas tenter de récupérer le fichier partiel."),
            ("Batterie vide en cours de route",
             "Brancher la batterie externe. Si fichier coupé, refaire le tronçon plus tard."),
            ("Route bloquée (manifestation, accident, travaux)",
             "Arrêter la journalisation, noter sur la fiche, prévoir une nouvelle sortie un autre jour."),
            ("Doute sur l'emplacement exact du point de départ ou d'arrivée",
             "Utiliser Google Maps en parallèle. Les coordonnées GPS exactes sont sur la fiche du tronçon."),
            ("Oubli de noter l'heure de départ ou d'arrivée",
             "L'heure est dans le fichier GPX (premier et dernier point). Pas grave — noter ce qui est encore mémorisable."),
        ],
        largeurs_cm=[6.5, 9.5],
        premiere_col_grise=True,
    )

    saut_de_page(doc)

    # ----- 11. Récapitulatif final -----
    ajouter_h1(doc, "11. Récapitulatif final à compléter")
    ajouter_paragraphe(doc, "À remplir une fois rentré, avant de transmettre les fichiers :")
    doc.add_paragraph()

    ajouter_tableau(
        doc,
        entetes=["Champ", "À remplir"],
        lignes=[
            ("Date de la sortie", "____ / ____ / ________"),
            ("Nom du conducteur", "________________________________"),
            ("Nom de l'observateur (si présent)", "________________________________"),
            ("Véhicule utilisé", "________________________________"),
            ("Smartphone utilisé (marque/modèle)", "________________________________"),
            ("Application GPS utilisée", "________________________________"),
            ("Heure de début de la sortie", "____ h ____ min"),
            ("Heure de fin de la sortie", "____ h ____ min"),
            ("Météo générale du jour", "________________________________"),
            ("Nombre de tronçons réalisés", "____ / 6"),
            ("Nombre de fichiers GPX exploitables", "____ / 6"),
        ],
        largeurs_cm=[8.0, 8.0],
        premiere_col_grise=True,
    )

    doc.add_paragraph()
    ajouter_paragraphe(doc, "Remarques générales sur la sortie :", gras=True)
    for _ in range(5):
        ajouter_paragraphe(doc, "_" * 75, taille=11, espace_apres=2)

    doc.add_paragraph()
    encadre_info(doc, "Une fois ce document complété", [
        "1. Rassembler les 6 fichiers GPX dans un dossier.",
        "2. Renommer le dossier au format : paa_releves_terrain_AAAAMMJJ/",
        "3. Joindre une copie scannée ou photographiée de cette procédure remplie.",
        "4. Transmettre le tout au responsable FLUIDIS pour import en P5.",
    ])

    # Sauvegarde
    sortie = "Procedure_Collecte_GPX_Terrain_FLUIDIS.docx"
    doc.save(sortie)
    print(f"OK Document genere : {sortie}")
    return sortie


if __name__ == "__main__":
    generer()
